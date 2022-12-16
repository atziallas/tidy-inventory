from io import BytesIO, StringIO

from barcode import generate
from docx import Document
from sequences import get_next_value

from lib.barcode import generate_ean13
from lib.writer import ImageWriter
from tidy.settings import SEQUENCE_NAME
from django.contrib import messages

ACTION_FAIL = 'fail'

def generate_barcode(modeladmin, request, queryset):
    for thing in queryset:
        if thing.barcode is None:
            thing.barcode = generate_ean13(get_next_value(SEQUENCE_NAME))
            thing.save()
    modeladmin.message_user(request, "Barcodes generated. Objects tagged as not labeled yet.")


def unlabel(modeladmin, request, queryset):
    for thing in queryset:
        thing.labeled = False
        thing.save()
    modeladmin.message_user(request, "Barcodes removed")


def print_barcode(modeladmin, request, queryset):
    idx = 0
    with open('core/templates/labels/MR183.docx', 'rb') as f:
        source_stream = BytesIO(f.read())
    doc = Document(source_stream)
    source_stream.close()
    # doc = Document('core/templates/labels/MR183.docx')

    # the template text and barcode cells pattern is a bit convoluted
    # so for now I use just this to easily port to another template
    # the sticker labels template, doesnt have  
    barcode_pos = [
        [0, 0], [0, 4], [0, 8],
        [2, 2], [2, 6], [2, 10],

        [5, 0], [5, 4], [5, 8],
        [7, 2], [7, 6], [7, 10],

        [10, 0], [10, 4], [10, 8],
        [12, 2], [12, 6], [12, 10],

        [15, 0], [15, 4], [15, 8],
        [17, 2], [17, 6], [17, 10],

        [20, 0], [20, 4], [20, 8],
        [22, 2], [22, 6], [22, 10]
    ]

    text_pos = [
        [1, 0], [1, 4], [1, 8],
        [4, 2], [4, 6], [4, 10],

        [6, 0], [6, 4], [6, 8],
        [9, 2], [9, 6], [9, 10],

        [11, 0], [11, 4], [11, 8],
        [14, 2], [14, 6], [14, 10],

        [16, 0], [16, 4], [16, 8],
        [19, 2], [19, 6], [19, 10],

        [21, 0], [21, 4], [21, 8],
        [24, 2], [24, 6], [24, 10],
    ]

    for thing in queryset:
        fp = BytesIO()

        if not thing.barcode:
            modeladmin.message_user(request, "%s has no barcode" % thing, messages.ERROR)
            return ACTION_FAIL

        generate('EAN13', thing.barcode, writer=ImageWriter(), output=fp,
                 writer_options={
                     'font_size': 29,
                     'text_distance': 1.5,
                     'quiet_zone': 0,
                     'module_width': 0.66,
                     'module_height': 26.3
                 })

        docTable = doc.tables[0]
        add_barcode_image(barcode_pos, docTable, fp, idx)
        add_barcode_text(docTable, idx, text_pos, thing)
        thing.labeled = True
        thing.save()
        idx += 1

    # doc.save('core/templates/labels/MR183filled.docx')
    target_stream = BytesIO()
    doc.save(target_stream)
    modeladmin.message_user(request, "Word file saved succesfuly")
    return target_stream


def add_barcode_image(barcode_pos, docTable, fp, idx):
    barcode_paragraph = docTable.rows[barcode_pos[idx][0]].cells[barcode_pos[idx][1]].paragraphs[0]
    for run in barcode_paragraph.runs:
        run.text = ""
    barcode_run = barcode_paragraph.runs[0]
    barcode_run.text = ""
    barcode_run.add_picture(fp, 731100)


def add_barcode_text(docTable, idx, text_pos, thing):
    text_paragraph = docTable.rows[text_pos[idx][0]].cells[text_pos[idx][1]].paragraphs[0]
    for run in text_paragraph.runs:
        run.text = ""
    text_paragraph.runs[0].text = thing.name
